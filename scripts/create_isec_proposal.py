from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path("entregables/Propuesta_comercial_ISEC.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
PALE = "F3F6F9"
GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "1A1A1A"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.78)
sec.right_margin = Inches(0.78)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, before, after in [("Heading 1",16,16,7),("Heading 2",13,12,5),("Heading 3",11.5,8,4)]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLUE)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx]); margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def header_row(row):
    for cell in row.cells:
        shade(cell, BLUE)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.color.rgb = RGBColor.from_string(WHITE); r.font.size = Pt(9)

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p

def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.20)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p

# Running header/footer
hp = sec.header.paragraphs[0]
hp.text = "ISEC S.A.S.  |  Propuesta comercial"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(GRAY)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Propuesta confidencial · Vigencia: 15 días calendario")
run.font.size = Pt(8); run.font.color.rgb = RGBColor.from_string(GRAY)

# Cover / proposal centerpiece
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(48); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROPUESTA COMERCIAL")
r.font.name="Calibri"; r.font.size=Pt(25); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
p2 = doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(5)
r=p2.add_run("Ajustes y optimización del sitio web ISEC"); r.font.size=Pt(16); r.font.bold=True
p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER; p3.paragraph_format.space_after=Pt(28)
r=p3.add_run("Magento Open Source 2.4.7-p2"); r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(GRAY)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [("Cliente","ISEC S.A.S."),("Sitio web","https://isec.com.co"),("Tarifa","$120.000 COP por hora más IVA"),("Fecha","28 de agosto de 2026")]
for i,(a,b) in enumerate(data):
    meta.cell(i,0).text=a; meta.cell(i,1).text=b
    shade(meta.cell(i,0),LIGHT_BLUE)
    meta.cell(i,0).paragraphs[0].runs[0].bold=True
set_table_geometry(meta,[2300,7060])

doc.add_paragraph()
lead=doc.add_paragraph(); lead.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=lead.add_run("Valor propuesto: $24.000.000 COP más IVA")
r.font.size=Pt(15); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Bolsa de 200 horas · Dedicación promedio de 4 horas diarias · Duración estimada de 10 a 12 semanas").italic=True

doc.add_page_break()
doc.add_heading("1. Presentación", level=1)
doc.add_paragraph("De acuerdo con la reunión realizada y el checklist de inconvenientes compartido, presentamos la propuesta para analizar e implementar los ajustes visuales, funcionales y responsive solicitados para las páginas de Distribuidores y Proyectos del sitio web de ISEC.")
doc.add_paragraph("La plataforma se encuentra implementada en Magento Open Source 2.4.7-p2 y alojada en AWS. Inicialmente, los cambios continuarán publicándose mediante el procedimiento disponible actualmente.")

doc.add_heading("2. Objetivo", level=1)
doc.add_paragraph("Realizar los ajustes relacionados en el checklist suministrado, procurando conservar la estabilidad de las funcionalidades existentes y documentando los cambios implementados.")

doc.add_heading("3. Alcance y estimación", level=1)
items = [
(1,"Análisis página web","12–20 h","$1.440.000–$2.400.000"),
(2,"Mejorar la visual de la recepción de los formularios","6–12 h","$720.000–$1.440.000"),
(3,"Revisión íconos social media debajo de logo de ISEC en el home y pre home (responsive también)","4–8 h","$480.000–$960.000"),
(4,"Supermenú responsive","12–20 h","$1.440.000–$2.400.000"),
(5,"Categorías de productos","10–18 h","$1.200.000–$2.160.000"),
(6,"Plantilla de visualización productos","20–32 h","$2.400.000–$3.840.000"),
(7,"reCAPTCHA","3–6 h","$360.000–$720.000"),
(8,"Formulario creación de cuenta","14–26 h","$1.680.000–$3.120.000"),
(9,"Ajuste supermenú","10–18 h","$1.200.000–$2.160.000"),
(10,"Formulario contáctenos","6–12 h","$720.000–$1.440.000"),
(11,"Base de datos","8–16 h","$960.000–$1.920.000"),
(12,"Quitar carrito del supermenú","2–4 h","$240.000–$480.000"),
(13,"Logo ISEC","2–4 h","$240.000–$480.000"),
(14,"Agregar ícono casa","2–3 h","$240.000–$360.000"),
(15,"Recarga del prehome","6–12 h","$720.000–$1.440.000"),
(16,"Plantillas","16–30 h","$1.920.000–$3.600.000"),
(17,"Ajuste copyright","2–3 h","$240.000–$360.000"),
(18,"Capacitaciones","8–12 h","$960.000–$1.440.000"),
]
t=doc.add_table(rows=1,cols=4)
t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,v in enumerate(["ÍTEM","AJUSTE","ESTIMACIÓN","VALOR ESTIMADO"]): t.cell(0,i).text=v
header_row(t.rows[0])
for idx,title,hours,value in items:
    cells=t.add_row().cells
    cells[0].text=str(idx); cells[1].text=title; cells[2].text=hours; cells[3].text=value
    cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for c in cells:
        for p in c.paragraphs:
            p.paragraph_format.space_after=Pt(0)
            for rr in p.runs: rr.font.size=Pt(8.2)
total=t.add_row().cells
total[0].merge(total[1]); total[0].text="TOTAL ESTIMADO"; total[2].text="143–256 h"; total[3].text="$17.160.000–$30.720.000"
for c in total:
    shade(c,LIGHT_BLUE)
    for p in c.paragraphs:
        for rr in p.runs: rr.bold=True; rr.font.size=Pt(8.5)
set_table_geometry(t,[650,4490,1350,2870])
t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

doc.add_paragraph("Las cifras anteriores son estimaciones por actividad. Algunos ajustes comparten componentes y podrán ejecutarse conjuntamente. El consumo real se controlará contra la bolsa contratada.")

doc.add_heading("4. Propuesta económica", level=1)
summary=doc.add_table(rows=6,cols=2); summary.style="Table Grid"
econ=[("Bolsa de trabajo","200 horas"),("Dedicación promedio","4 horas diarias, de lunes a viernes"),("Tarifa por hora","$120.000 COP más IVA"),("Valor propuesto","$24.000.000 COP más IVA"),("Duración estimada","10–12 semanas calendario"),("Horas adicionales","$120.000 COP más IVA, previa autorización")]
for i,(a,b) in enumerate(econ):
    summary.cell(i,0).text=a; summary.cell(i,1).text=b; shade(summary.cell(i,0),PALE); summary.cell(i,0).paragraphs[0].runs[0].bold=True
set_table_geometry(summary,[3000,6360])
doc.add_paragraph("La estimación podrá revisarse después de completar Análisis página web, debido a que inicialmente no se cuenta con acceso al código ni con información completa sobre la construcción del proyecto. No se ejecutarán horas adicionales sin autorización del cliente.")
doc.add_paragraph("Las horas ejecutadas se registrarán y compartirán con el cliente mediante un archivo de Excel.")

doc.add_heading("5. Propuesta comercial y metodología / forma de trabajo", level=1)
for txt in [
"Revisión del funcionamiento actual y análisis del código relacionado con el checklist.",
"Priorización de los ítems con el cliente.",
"Desarrollo de los ajustes por bloques.",
"Presentación de los cambios para revisión.",
"Aplicación de observaciones.",
"Pruebas visuales, funcionales y responsive.",
"Publicación de los cambios aprobados mediante el procedimiento disponible actualmente.",
"Entrega de documentación y capacitación."]:
    add_number(txt)

doc.add_heading("6. Análisis y documentación de los cambios", level=1)
doc.add_paragraph("El análisis inicial se enfocará exclusivamente en los componentes necesarios para desarrollar el checklist e incluirá:")
for txt in [
"Revisión general de la estructura de carpetas.",
"Identificación del tema, módulos y plantillas relacionados con los ajustes.",
"Identificación de los archivos que deberán modificarse.",
"Revisión general del funcionamiento de formularios, menús, productos, usuarios y prehome.",
"Identificación de dependencias o personalizaciones que puedan afectar el desarrollo.",
"Registro de riesgos o bloqueos directamente relacionados con el alcance."]:
    add_bullet(txt)
doc.add_paragraph("La documentación técnica de los cambios se almacenará junto al código intervenido, preferiblemente en un archivo README.md general o dentro del módulo o tema correspondiente. Incluirá la descripción del cambio, los componentes intervenidos, las configuraciones requeridas y las consideraciones básicas para su validación.")

doc.add_heading("7. Auditoría documentada paso a paso y acompañamiento mediante cápsulas en video", level=1)
doc.add_paragraph("Los hallazgos relacionados con el checklist y los cambios implementados quedarán documentados progresivamente. Las cápsulas en video se utilizarán para explicar los procesos que requieran demostración práctica, tales como la administración de categorías y productos, el uso de formularios, las plantillas modificadas o los cambios administrativos que posteriormente pueda realizar el cliente.")
doc.add_paragraph("La documentación y los videos se limitarán a los componentes incluidos en el checklist y efectivamente intervenidos.")

doc.add_heading("8. Tarifa por horas de desarrollo para ajustes técnicos específicos", level=1)
doc.add_paragraph("Los ajustes técnicos adicionales que no se encuentren dentro del checklist tendrán una tarifa de $120.000 COP por hora más IVA. Antes de iniciar una actividad adicional se informará su descripción, estimación de horas y valor estimado. La ejecución requerirá aprobación previa del cliente.")

doc.add_heading("9. Entregables", level=1)
for txt in [
"Ajustes aprobados correspondientes al checklist.",
"Archivo de Excel con el registro de horas y actividades.",
"Documentación técnica de los cambios en archivos README.md.",
"Relación de archivos modificados.",
"Evidencias básicas de las pruebas realizadas.",
"Cápsulas de video correspondientes a las capacitaciones.",
"Sesión de entrega y socialización."]:
    add_bullet(txt)

doc.add_heading("10. Recomendaciones opcionales — Nice to have", level=1)
doc.add_paragraph("Durante el desarrollo podrán identificarse oportunidades de mejora que no forman parte del alcance principal, como:")
for txt in [
"Crear un repositorio en GitHub, GitLab, Bitbucket u otro gestor de versionamiento.",
"Organizar un flujo controlado para la publicación de cambios.",
"Crear un ambiente independiente de pruebas.",
"Revisar la actualización de Magento y sus parches de seguridad.",
"Mejorar la documentación técnica general del proyecto."]:
    add_bullet(txt)
doc.add_paragraph("Estas actividades se presentarán únicamente como recomendaciones y no se ejecutarán ni consumirán horas sin autorización previa del cliente.")

doc.add_heading("11. Supuestos", level=1)
doc.add_paragraph("Para ejecutar los ajustes, el cliente facilitará los accesos e información necesarios, incluyendo código del sitio, panel administrativo de Magento, archivos, base de datos cuando corresponda, contenidos y una persona responsable de revisar y aprobar los cambios.")

doc.add_heading("12. Forma de pago propuesta", level=1)
pay=doc.add_table(rows=5,cols=3); pay.style="Table Grid"
for i,v in enumerate(["MOMENTO","PORCENTAJE","VALOR"]): pay.cell(0,i).text=v
header_row(pay.rows[0])
payments=[("Inicio","40 %","$9.600.000 más IVA"),("Entrega del primer bloque","30 %","$7.200.000 más IVA"),("Entrega del segundo bloque","20 %","$4.800.000 más IVA"),("Entrega final","10 %","$2.400.000 más IVA")]
for i,row in enumerate(payments,1):
    for j,v in enumerate(row): pay.cell(i,j).text=v
    pay.cell(i,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    pay.cell(i,2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
set_table_geometry(pay,[4200,1800,3360])

doc.add_heading("13. Vigencia y aceptación", level=1)
doc.add_paragraph("La propuesta tiene una vigencia de 15 días calendario. El inicio del proyecto estará sujeto a la aprobación escrita, el pago inicial y la entrega de los accesos necesarios.")
doc.add_paragraph("\nNombre: ____________________________________________")
doc.add_paragraph("Cargo: _____________________________________________")
doc.add_paragraph("Firma: _____________________________________________")
doc.add_paragraph("Fecha: _____________________________________________")

# Keep rows readable and repeat header on table pages.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True

doc.core_properties.title = "Propuesta comercial - Ajustes sitio web ISEC"
doc.core_properties.subject = "Ajustes y optimización del sitio web ISEC"
doc.core_properties.author = "Propuesta comercial"
doc.save(OUT)
print(OUT.resolve())
